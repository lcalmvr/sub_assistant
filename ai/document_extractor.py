"""
Document text extraction module for insurance policy documents.
Supports PDF and DOCX formats.
"""

import os
from typing import Optional


class DocumentExtractionError(Exception):
    """Raised when document text extraction fails."""
    pass


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using pypdf.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text content

    Raises:
        DocumentExtractionError: If extraction fails
    """
    try:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        text = "\n\n".join(text_parts)

        # If we got very little text, the PDF might be scanned/image-based
        # Try unstructured as a fallback
        if len(text.strip()) < 100:
            fallback_text = _extract_with_unstructured(file_path)
            if len(fallback_text) > len(text):
                return fallback_text

        return text

    except ImportError:
        raise DocumentExtractionError("pypdf library not installed")
    except Exception as e:
        raise DocumentExtractionError(f"Failed to extract text from PDF: {e}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content

    Raises:
        DocumentExtractionError: If extraction fails
    """
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)

    except ImportError:
        raise DocumentExtractionError("python-docx library not installed")
    except Exception as e:
        raise DocumentExtractionError(f"Failed to extract text from DOCX: {e}")


def _extract_with_unstructured(file_path: str) -> str:
    """
    Fallback extraction using unstructured library.
    Useful for scanned PDFs or complex documents.

    Args:
        file_path: Path to the document

    Returns:
        Extracted text content
    """
    try:
        from unstructured.partition.auto import partition

        elements = partition(filename=file_path)
        text_parts = [str(el) for el in elements if str(el).strip()]
        return "\n".join(text_parts)

    except ImportError:
        return ""
    except Exception:
        return ""


def extract_text_from_document(file_path: str) -> str:
    """
    Auto-detect file type and extract text.

    Args:
        file_path: Path to the document (PDF or DOCX)

    Returns:
        Extracted text content

    Raises:
        DocumentExtractionError: If file type is unsupported or extraction fails
    """
    if not os.path.exists(file_path):
        raise DocumentExtractionError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        if ext == '.doc':
            # .doc files may not work with python-docx
            # Try unstructured first
            text = _extract_with_unstructured(file_path)
            if not text:
                raise DocumentExtractionError(
                    ".doc format requires conversion. Please save as .docx or PDF."
                )
        else:
            text = extract_text_from_docx(file_path)
    else:
        raise DocumentExtractionError(
            f"Unsupported file type: {ext}. Supported types: PDF, DOCX"
        )

    if not text.strip():
        raise DocumentExtractionError(
            "No text could be extracted from the document. "
            "The document may be image-based or empty."
        )

    return text


def extract_coverage_sections(text: str, max_length: int = 15000) -> str:
    """
    Pre-process text to focus on coverage-relevant sections.
    Useful for large documents to reduce token usage.

    Args:
        text: Full document text
        max_length: Maximum output length

    Returns:
        Filtered text focusing on coverage information
    """
    # Keywords that indicate coverage-relevant sections
    coverage_keywords = [
        "coverage", "sublimit", "limit", "deductible", "retention",
        "insuring agreement", "schedule", "endorsement",
        "aggregate", "occurrence", "waiting period",
        "ransomware", "extortion", "business interruption",
        "network security", "privacy", "regulatory",
        "social engineering", "fraud", "breach response",
        "media liability", "cyber", "technology"
    ]

    lines = text.split('\n')
    relevant_lines = []
    include_next = 0  # Include N lines after a keyword match

    for line in lines:
        line_lower = line.lower()

        # Check if line contains any coverage keywords
        has_keyword = any(kw in line_lower for kw in coverage_keywords)

        # Check if line contains dollar amounts (likely limit info)
        has_amount = any(c in line for c in ['$', 'USD']) or \
                     any(suffix in line.upper() for suffix in ['K', 'M', '000'])

        if has_keyword or has_amount or include_next > 0:
            relevant_lines.append(line)
            if has_keyword:
                include_next = 3  # Include next 3 lines for context
            else:
                include_next = max(0, include_next - 1)

    filtered_text = '\n'.join(relevant_lines)

    # If we filtered too aggressively, return original (truncated)
    if len(filtered_text) < 500:
        return text[:max_length]

    # Truncate if still too long
    if len(filtered_text) > max_length:
        return filtered_text[:max_length] + "\n\n[... truncated ...]"

    return filtered_text
